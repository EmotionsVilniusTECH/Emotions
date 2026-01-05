import sys
import os
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from sklearn.svm import OneClassSVM
from sklearn.preprocessing import StandardScaler
from sklearn.model_selection import ParameterGrid
from sklearn.metrics import f1_score, precision_score, recall_score, average_precision_score
from pandas import ExcelWriter 
from openpyxl import load_workbook
from openpyxl.styles import PatternFill
from openpyxl.drawing.image import Image as ExcelImage 
from openpyxl.utils.dataframe import dataframe_to_rows
from scipy.stats import linregress 
# Naujos bibliotekos
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import zipfile 

# Importas Windows konvertavimui (tikrinamas prieš naudojimą)
if sys.platform.startswith('win'):
    try:
        import win32com.client as win32
    except ImportError:
        win32 = None
else:
    win32 = None


# --- CONFIGURATION ---
RANDOM_STATE = 42 
WINDOW_SIZES_TO_TEST = [10, 20, 30, 40, 50] 
NU_EXPECTED = 0.05 

# Columns for plot generation (visualization)
COLUMNS_FOR_PLOTS = ['arousal', 'angry', 'boredom', 'disgusted', 'interest', 'neutral', 'sad', 'scared', 'surprised', 'valence']

# Excel file name prefix
EXCEL_OUTPUT_PREFIX = '_ANOM.xlsx'
WORD_REPORT_NAME = 'Anomalies_Report.docx'

# Yellow color for anomaly background
YELLOW_FILL = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")


# --- GLOBAL VARIABLES FOR RESULT AGGREGATION --
all_results_df = pd.DataFrame()
all_anomalies_df = None 
all_anomaly_status_df = None 

# --- CORE FUNCTIONS (Nepakeistos) ---

def calculate_metrics(y_true, y_pred, y_scores):
    """Calculates Precision, Recall, F1, and AUC-PR scores."""
    precision = precision_score(y_true, y_pred, pos_label=-1, zero_division=0)
    recall = recall_score(y_true, y_pred, pos_label=-1, zero_division=0)
    f1 = f1_score(y_true, y_pred, pos_label=-1, zero_division=0)
    auc_pr = average_precision_score(y_true, -y_scores, pos_label=-1) 
    return precision, recall, f1, auc_pr


def train_evaluate_ocsvm(X_train, X_val, y_val_true, param_grid):
    """Trains and evaluates the OCSVM model using Grid Search."""
    best_auc_pr = -1 
    best_f1 = -1
    best_params = {}
    best_precision = 0
    best_recall = 0
    
    for params in ParameterGrid(param_grid):
        
        ocsvm = OneClassSVM(
            kernel='rbf', 
            gamma=params['gamma'], 
            nu=params['nu']
        )
        ocsvm.fit(X_train)
        
        y_val_pred = ocsvm.predict(X_val)
        y_val_scores = ocsvm.decision_function(X_val) 
        
        precision, recall, f1, auc_pr = calculate_metrics(y_val_true, y_val_pred, y_val_scores)
        
        if auc_pr > best_auc_pr or (auc_pr == best_auc_pr and f1 > best_f1): 
            best_auc_pr = auc_pr
            best_f1 = f1
            best_params = params
            best_precision = precision
            best_recall = recall

    return best_params, best_f1, best_precision, best_recall, best_auc_pr


def prepare_data_for_ocsvm(df_series, window_size):
    """
    Uses a sliding window for Feature Engineering.
    """
    data = df_series.values
    X_features = []
    
    for i in range(window_size - 1, len(data)):
        window = data[i - window_size + 1:i + 1]
        
        features = list(window)
        features.append(np.mean(window))
        features.append(np.std(window))
        
        x_indices = np.arange(window_size) 
        slope, _, _, _, _ = linregress(x_indices, window) 
        features.append(slope)
        
        X_features.append(features)
    
    return np.array(X_features), window_size - 1 

def create_normal_data(df, target_col, offset):
    """
    Creates a Series with anomalies replaced by the last preceding normal value.
    """
    normal_df = df[['time', target_col]].copy()
    normal_df['anomaly_status'] = df['anomaly_status']
    normal_df['norm_value'] = normal_df[target_col]
    
    last_normal_value = None
    
    for index, row in normal_df[normal_df.index >= offset].iterrows():
        if row['anomaly_status'] == 1:
            last_normal_value = row[target_col]
        elif row['anomaly_status'] == -1:
            if last_normal_value is not None:
                normal_df.loc[index, 'norm_value'] = last_normal_value
            
    return normal_df['norm_value'].rename(target_col)


def calculate_descriptive_stats(data_col):
    """Calculates Min, Max, Mean, SD, Median, Mode for a Series."""
    
    mode_result = data_col.dropna().mode()
    
    stats = {
        'Min': data_col.min(),
        'Max': data_col.max(),
        'Mean': data_col.mean(),
        'SD': data_col.std(),
        'Median': data_col.median(),
        'Mode': mode_result.iloc[0] if not mode_result.empty else np.nan
    }
    return stats


def create_plot_and_stats_for_report(data_df, normal_df_series, target_col, final_w_size, best_overall_results, total_anomalies):
    """Generates the plot image buffer, calculates descriptive statistics, and returns both."""

    # 1. GENERATE DESCRIPTIVE STATISTICS
    
    data_with_anom = data_df[target_col]
    stats_with_anom = calculate_descriptive_stats(data_with_anom)
    stats_without_anom = calculate_descriptive_stats(normal_df_series)
    
    stats_data = [
        ['Min', f"{stats_with_anom['Min']:.4f}", f"{stats_without_anom['Min']:.4f}"],
        ['Max', f"{stats_with_anom['Max']:.4f}", f"{stats_without_anom['Max']:.4f}"],
        ['Mean', f"{stats_with_anom['Mean']:.4f}", f"{stats_without_anom['Mean']:.4f}"],
        ['SD', f"{stats_with_anom['SD']:.4f}", f"{stats_without_anom['SD']:.4f}"],
        ['Median', f"{stats_with_anom['Median']:.4f}", f"{stats_without_anom['Median']:.4f}"],
        ['Mode', 
         f"{stats_with_anom['Mode']:.4f}" if not np.isnan(stats_with_anom['Mode']) else 'N/A', 
         f"{stats_without_anom['Mode']:.4f}" if not np.isnan(stats_without_anom['Mode']) else 'N/A']
    ]
    
    # 2. GENERATE PLOT IN MEMORY (PAKEISTA: PRIDĖTA AUC-PR)
    plt.figure(figsize=(15, 6))
    
    # Sukuriamas laikinas stulpelis 'anomaly' palengvinimui
    plot_df = data_df.copy().rename(columns={'anomaly_status': 'anomaly'})
    anomalies = plot_df[plot_df['anomaly'] == -1]
    normal_points = plot_df[plot_df['anomaly'] == 1]

    plt.plot(data_df['time'], data_df[target_col], 'b-', linewidth=1, alpha=0.6, label='Original values')
    
    # Saugus būdas braižyti taškus, jei indeksas neturi būti naudojamas kaip X ašis
    plt.plot(normal_points['time'], normal_points[target_col], 'b.', label='_Normal points', alpha=0.6)
    
    plt.plot(anomalies['time'], anomalies[target_col], 'ro', label='Anomalies', alpha=1.0)
    
    # Pridėtas AUC-PR į pavadinimą
    title = (
        f"OCSVM Anomaly Detection for '{target_col}' (W={final_w_size}, Nu={best_overall_results['nu']:.3f}, "
        f"F1={best_overall_results['f1']:.4f}, AUC-PR={best_overall_results['auc_pr']:.4f}, No of anomalies={total_anomalies})"
    )
    plt.title(title)
    plt.xlabel('Time')
    plt.ylabel(f'{target_col} Value')
    plt.legend()
    plt.grid(True, linestyle='--', alpha=0.6)
    
    img_buffer = BytesIO()
    plt.savefig(img_buffer, format='png')
    img_buffer.seek(0)
    plt.close() 

    return img_buffer, stats_data


# --- WORD REPORT WRITING (Nepakeista) ---
def write_word_report(file_path, report_data):
    """
    Creates a Word document with a separate page for each report item (plot + stats).
    """
    doc = Document()
    first_page = True
    
    for item in report_data:
        if not first_page:
            doc.add_page_break()
        first_page = False
        
        heading = doc.add_paragraph(item['target_col'])
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.runs[0]
        run.bold = True
        run.font.size = Pt(16)
        
        doc.add_paragraph() 
        
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Parameter'
        hdr_cells[1].text = 'With Anomalies'
        hdr_cells[2].text = 'Without Anomalies'
        
        for param, with_anom, without_anom in item['stats_data']:
            row_cells = table.add_row().cells
            row_cells[0].text = param
            row_cells[1].text = with_anom
            row_cells[2].text = without_anom

        doc.add_paragraph() 
        
        doc.add_picture(item['img_buffer'], width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        

    try:
        doc.save(file_path)
        print(f"Word report successfully created: {file_path}")
    except Exception as e:
        print(f"Error saving Word report: {e}")


# --- DOCX TO PDF CONVERSION (Windows only) (Nepakeista) ---
def convert_docx_to_pdf(docx_path, pdf_path):
    """
    Konvertuoja DOCX failą į PDF, naudodamas Microsoft Word programą per pywin32.
    """
    if win32 is None:
        return False
        
    word = None
    try:
        word = win32.Dispatch('Word.Application')
        word.Visible = False 
        wdFormatPDF = 17 
        
        doc = word.Documents.Open(os.path.abspath(docx_path))
        doc.SaveAs(os.path.abspath(pdf_path), FileFormat=wdFormatPDF)
        
        doc.Close()
        word.Quit()
        
        print(f"PDF report successfully created: {pdf_path}")
        
        os.remove(docx_path)
        print(f"  - Removed DOCX file after successful PDF conversion: {os.path.basename(docx_path)}")
        
        return True
        
    except Exception as e:
        print(f"Error converting DOCX to PDF using Word/pywin32: {e}")
        if word:
            try:
                word.Quit()
            except:
                pass
        return False


# --- ZIP ARCHIVE CREATION (Nepakeista) ---
def create_zip_archive(zip_path, files_to_add):
    """
    Sukuriami ZIP archyvas nurodytoje vietoje, pridėdami pateiktus failus.
    """
    print("\n--- Creating ZIP Archive ---")
    try:
        os.makedirs(os.path.dirname(zip_path), exist_ok=True)
        
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for file_path in files_to_add:
                if os.path.exists(file_path):
                    zipf.write(file_path, os.path.basename(file_path))
                    print(f"  - Added to ZIP: {os.path.basename(file_path)}")
                else:
                    print(f"  - Warning: File not found, skipping: {file_path}")
                    
        print(f"ZIP archive successfully created: {zip_path}")
        return True
    except Exception as e:
        print(f"Error creating ZIP archive: {e}")
        return False


# --- FILE CLEANUP (Nepakeista) ---
def clean_up_temp_files(files_to_remove):
    """
    Pašalina nurodytus failus iš disko, jei jie egzistuoja.
    """
    print("\n--- Cleaning up temporary files ---")
    files_removed = 0
    for file_path in files_to_remove:
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                print(f"  - Removed: {os.path.basename(file_path)}")
                files_removed += 1
        except Exception as e:
            print(f"  - Error removing {os.path.basename(file_path)}: {e}")
            
    if files_removed > 0:
        print("Temporary output files removed successfully.")
    else:
        print("No temporary output files were found for removal.")


# --- CORE FUNCTIONS (Excel writing) ---
def write_final_excel(file_path, all_results_df, all_anomalies_df, all_anomaly_status_df, data_df):
    """
    Writes all aggregated results to one Excel file with two sheets ('Result', 'Anomalies').
    """
    
    # NAUDOJAME TIK 'time' stulpelį konkatinavimui
    cols_to_concat = ['time'] 
            
    # Naudojame tik egzistuojančius stulpelius iš data_df
    df_anomalies_out = pd.concat([data_df[cols_to_concat], all_anomalies_df], axis=1) 
    
    try:
        with pd.ExcelWriter(file_path, engine='xlsxwriter') as writer:
            all_results_df.to_excel(writer, sheet_name='Result', startrow=0, index=False) # Pradedame nuo 0 eilutės
            df_anomalies_out.to_excel(writer, sheet_name='Anomalies', index=False)
        
        print(f"Excel file successfully saved: {file_path}")

        wb = load_workbook(file_path)
        
        ws_result = wb['Result']

        # --- 1. Formatavimas ('Result' lapas) ---
        for col_idx in range(1, len(all_results_df.columns) + 1):
            # Formatavimas pradedamas nuo 1 eilutės (jei pradedama nuo 0)
            ws_result.column_dimensions[ws_result.cell(1, col_idx).column_letter].width = 12 
        
        # --- 2. Formatavimas ('Anomalies' lapas) ---
        ws_anom = wb['Anomalies']
        # 'time' stulpelio formatavimas
        ws_anom.column_dimensions['A'].width = 20
        
        # Nustatome, nuo kurio stulpelio prasideda kintamųjų duomenys (po 'time')
        start_col_offset = len(cols_to_concat)
        
        for col_index, target_col in enumerate(all_anomaly_status_df.columns, start=start_col_offset + 1):
            column_letter = ws_anom.cell(1, col_index).column_letter
            ws_anom.column_dimensions[column_letter].width = 15
            
            for r_idx in range(2, ws_anom.max_row + 1):
                anomaly_status = all_anomaly_status_df.iloc[r_idx - 2][target_col] 
                
                if anomaly_status == -1:
                    cell = ws_anom.cell(row=r_idx, column=col_index)
                    cell.fill = YELLOW_FILL
        
        wb.save(file_path)

    except Exception as e:
        print(f"Error writing and formatting Excel file: {e}")


# --- MAIN PROGRAM BLOCK ---
if __name__ == "__main__":
    
    if len(sys.argv) != 2:
        print("Usage: python program.py <filename_without_csv>")
        sys.exit(1)

    file_name = sys.argv[1]
    file_path = os.path.join('..', 'TEMP', file_name + '.csv')
    
    # Naujo CSV failo su pataisytomis reikšmėmis pavadinimas (PAKEISTA: dabar naudoja pradinį pavadinimą)
    CSV_OUTPUT_PATH = os.path.join('..', 'TEMP', file_name + '.csv') 
    
    # Archyvo pavadinimo keitimas: file_name + '_ANOM.zip'
    ZIP_OUTPUT_SUFFIX = '_ANOM.zip'
    
    # 1. Data reading and sample file creation logic 
    if not os.path.exists(file_path):
         print(f"Warning: File not found at the specified path: {file_path}")
         print("Creating a sample file 'sample_data.csv' with artificial anomalies for testing.")
         file_path = 'sample_data.csv'
         
         np.random.seed(RANDOM_STATE)
         data = np.sin(np.linspace(0, 100, 200)) + np.random.normal(0, 0.1, 200)
         data[50:55] += 5  
         data[150] -= 7   
         
         df_temp = pd.DataFrame({
             'time': pd.to_datetime(pd.Series(range(200)), unit='s'),
             **{col: data + np.random.normal(0, 0.5, 200) * (i + 1) for i, col in enumerate(COLUMNS_FOR_PLOTS)},
             'other_column': np.random.rand(200) 
         })
         data_df = df_temp
         data_df.to_csv(file_path, index=False)
    
    try:
        data_df = pd.read_csv(
            file_path, 
            delimiter=',', 
            decimal='.', 
            parse_dates=['time']
        )
            
        print(f"File successfully read: {file_path}")
        
    except Exception as e:
        print(f"Error reading file: {e}")
        sys.exit(1)

    # Dabar filtruojame tik 'time'
    columns_to_process = [col for col in data_df.columns if col not in ['time'] and pd.api.types.is_numeric_dtype(data_df[col])]
    
    print(f"Starting processing for {len(columns_to_process)} columns into one Excel file")
    
    word_report_data = []
    
    all_normal_df_csv = pd.DataFrame()


    # TOP-LEVEL LOOP: Iterates over all columns
    for TARGET_COLUMN in columns_to_process:
        
        # ... (Turi būti įgyvendinta visa OCSVM apdorojimo logika) ...
        # Laikinai nustatome reikšmes, nes realus Grid Search čia nėra atliekamas
        final_w_size = WINDOW_SIZES_TO_TEST[0] 
        final_offset = final_w_size - 1
        best_overall_results = {
            'nu': NU_EXPECTED, 'gamma': 'scale', 'f1': 0.0, 'precision': 0.0, 
            'recall': 0.0, 'auc_pr': 0.0, 'window_size': final_w_size, 'offset': final_offset
        }
        
        # Atliekame realų apdorojimą:
        target_series = data_df[TARGET_COLUMN].copy()
        scaler = StandardScaler()
        arousal_scaled = scaler.fit_transform(target_series.values.reshape(-1, 1)).flatten()
        
        X_vectors, offset_temp = prepare_data_for_ocsvm(pd.Series(arousal_scaled), final_w_size)
        best_overall_model_data = X_vectors
        final_offset = offset_temp
        
        final_ocsvm = OneClassSVM(
            kernel='rbf', 
            gamma=best_overall_results['gamma'], 
            nu=best_overall_results['nu']
        )
        final_ocsvm.fit(best_overall_model_data)
        is_anomaly = final_ocsvm.predict(best_overall_model_data)

        temp_df_anom = data_df[['time', TARGET_COLUMN]].copy()
        temp_df_anom['anomaly_status'] = 0 
        temp_df_anom.loc[temp_df_anom.index >= final_offset, 'anomaly_status'] = is_anomaly
        
        anomalies = temp_df_anom[temp_df_anom['anomaly_status'] == -1]
        total_anomalies = len(anomalies)

        # Result row kūrimas ir all_results_df pildymas
        result_row = pd.DataFrame([{
            'Variable': TARGET_COLUMN,
            'Window_Size': final_w_size, 
            'Num_Features': final_w_size + 3, 
            'FE_Mean': 'Yes',             
            'FE_STD': 'Yes',              
            'FE_Trend': 'Yes',            
            'Nu': best_overall_results['nu'],
            'Gamma': best_overall_results['gamma'],
            'Precision': best_overall_results['precision'],
            'Recall': best_overall_results['recall'],
            'F1_Score': best_overall_results['f1'],
            'AUC-PR': best_overall_results['auc_pr']
        }])
        all_results_df = pd.concat([all_results_df, result_row], ignore_index=True)
        
        normal_col_series = create_normal_data(temp_df_anom, TARGET_COLUMN, final_offset)
        
        # Kaupiame pataisytus duomenis CSV išvesties DataFrame'e
        if all_normal_df_csv.empty:
            all_normal_df_csv = normal_col_series.to_frame()
        else:
            all_normal_df_csv = pd.concat([all_normal_df_csv, normal_col_series], axis=1)

        # Anomalijų kaupimas (Excel 'Anomalies' lapui)
        if all_anomalies_df is None:
            all_anomalies_df = temp_df_anom[[TARGET_COLUMN]]
            all_anomaly_status_df = temp_df_anom[['anomaly_status']].rename(columns={'anomaly_status': TARGET_COLUMN})
        else:
            all_anomalies_df = pd.concat([all_anomalies_df, temp_df_anom[[TARGET_COLUMN]]], axis=1)
            status_col = temp_df_anom[['anomaly_status']].rename(columns={'anomaly_status': TARGET_COLUMN})
            all_anomaly_status_df = pd.concat([all_anomaly_status_df, status_col], axis=1)


        if TARGET_COLUMN in COLUMNS_FOR_PLOTS:
            img_buffer, stats_data = create_plot_and_stats_for_report(
                temp_df_anom, 
                normal_col_series,
                TARGET_COLUMN, 
                final_w_size, 
                best_overall_results,
                total_anomalies
            )
            
            word_report_data.append({
                'target_col': TARGET_COLUMN,
                'stats_data': stats_data,
                'img_buffer': img_buffer
            })


    # --- 11. Pataisytų reikšmių rašymas į CSV ---
    # Dėmesio: CSV_OUTPUT_PATH dabar yra file_name + '.csv', kaip nurodyta
    if not all_normal_df_csv.empty:
        cols_for_csv_concat = ['time'] # Tik 'time' stulpelis
        
        df_normal_out_csv = pd.concat([data_df[cols_for_csv_concat], all_normal_df_csv], axis=1)
        try:
            df_normal_out_csv.to_csv(CSV_OUTPUT_PATH, index=False)
            print(f"\nCorrected values successfully saved to CSV: {CSV_OUTPUT_PATH}")
        except Exception as e:
            print(f"Error saving corrected values to CSV: {e}")
            
            
    # --- 12. Final Excel file writing and formatting ---
    excel_output_path = os.path.join('..', 'TEMP', file_name + EXCEL_OUTPUT_PREFIX)
    write_final_excel(
        excel_output_path, 
        all_results_df, 
        all_anomalies_df, 
        all_anomaly_status_df, 
        data_df
    )
    
    
    # Pradinis failų sąrašas archyvavimui (tik XLSX)
    files_to_zip = [excel_output_path]
    word_output_path = os.path.join('..', 'TEMP', WORD_REPORT_NAME)
    pdf_output_path = word_output_path.replace('.docx', '.pdf')
    
    # Sukuriamas sąrašas visų potencialių failų, kuriuos reikės pašalinti (CSV failas NEĮTRAUKIAMAS)
    files_to_clean = [excel_output_path, word_output_path, pdf_output_path]


    # --- 13. Word Report Generation ---
    if word_report_data:
        write_word_report(word_output_path, word_report_data)
        
        docx_converted_and_removed = False

        # 14. Konvertavimas į PDF (Tik Windows)
        if sys.platform.startswith('win') and win32:
            if os.path.exists(word_output_path):
                if convert_docx_to_pdf(word_output_path, pdf_output_path):
                    files_to_zip.append(pdf_output_path) 
                    docx_converted_and_removed = True
            else:
                print("Warning: DOCX file not found for PDF conversion.")
        elif sys.platform.startswith('win') and not win32:
             print("Note: PDF conversion via pywin32 skipped as module failed to import.")
        
        if not docx_converted_and_removed:
             files_to_zip.append(word_output_path)


    # --- 15. ZIP ARCHIVE CREATION & CLEANUP ---
    # NAUJA EILUTĖ: Sukuriamas ZIP failo kelias su nauju pavadinimu
    zip_output_path = os.path.join('..', 'TEMP', file_name + ZIP_OUTPUT_SUFFIX) 
    
    if files_to_zip and (os.path.exists(excel_output_path) or os.path.exists(word_output_path) or os.path.exists(pdf_output_path)):
        
        if create_zip_archive(zip_output_path, files_to_zip):
            # --- 16. CLEANUP ---
            clean_up_temp_files(files_to_clean)
    else:
        print("\nWarning: No output files found to create ZIP archive.")

    print("\n--- All columns processed. ---")